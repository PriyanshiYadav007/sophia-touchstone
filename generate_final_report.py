"""
Final BULLETPROOF Task Compilation Script for Bloom Valley Nursery.
Adds detailed analysis for accessibility, RWD, and storage as per rubric.
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = r"c:\Users\DELL\OneDrive\Desktop\Sophia.1web dev\webdev\bloom-valley-nursery\submissions\Final_Touchstone_Report.docx"
SCREENSHOT_DIR = r"c:\Users\DELL\OneDrive\Desktop\Sophia.1web dev\webdev\bloom-valley-nursery\final_screenshots"

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level)
    if level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_image(doc, filename, caption):
    path = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.5))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style.font.italic = True
    else:
        doc.add_paragraph(f"[Note: Image {filename} placeholder for final review]")

def main():
    doc = Document()
    
    # --- Title Page ---
    add_heading(doc, 'Final Touchstone Report', 0)
    p = doc.add_paragraph('\nBloom Valley Nursery Website Project\n\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Developed by: Priyanshi Yadav\nDate: February 11, 2026\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Section 1: Introduction & Design ---
    add_heading(doc, '1. Introduction and Design Phase', 1)
    doc.add_paragraph('Bloom Valley Nursery is a family-owned business specializing in plants, trees, and gardening supplies. The website aims to improve customer engagement and facilitate custom orders.')
    
    doc.add_paragraph('Design Rationale:')
    doc.add_paragraph('- Color Palette: Leafy Green (#4A7C59) and harvest gold accents were used to represent growth and nature.')
    doc.add_paragraph('- Typography: Playfair Display (Serif) for headers adds a touch of elegance, while Open Sans (Sans-Serif) ensures readability in body content.')
    doc.add_paragraph('- Logo: A tree-based emblem that reinforces the brand identity.')

    add_heading(doc, 'Desktop Wireframe', 2)
    add_image(doc, 'wireframe_desktop.png', 'Fig 1: Desktop Wireframe UI Layout')
    
    add_heading(doc, 'Mobile Wireframe', 2)
    add_image(doc, 'wireframe_mobile.png', 'Fig 2: Mobile Responsive Wireframe Layout')
    doc.add_page_break()

    # --- Section 2: Website Structure ---
    add_heading(doc, '2. Website Structure and HTML Components', 1)
    doc.add_paragraph('The website is built with a semantic HTML5 structure for better accessibility and SEO.')
    
    doc.add_paragraph('Component Descriptions:')
    doc.add_paragraph('- Navigation Bar: Persistent header allows quick access to Home, Gallery, About, and Community pages.')
    doc.add_paragraph('- Homepage: Features a Hero section with a CTA, Featured Products, and a Perks section.')
    doc.add_paragraph('- Gallery: Organized in a 3x3 grid using CSS Flexbox/Grid for a professional product showcase.')
    doc.add_paragraph('- About Us: Contains a business hours table and a comprehensive contact form for custom orders.')
    doc.add_paragraph('- Community Corner: A custom page highlighting customer testimonials and local events.')

    add_heading(doc, 'Homepage Rendering', 2)
    add_image(doc, '01_homepage_full.png', 'Fig 3: Live Homepage Screenshot')
    
    add_heading(doc, 'Gallery Rendering', 2)
    add_image(doc, '03_gallery_grid.png', 'Fig 4: Gallery Page Screenshot')
    doc.add_page_break()

    # --- Section 3: Styling & Accessibility ---
    add_heading(doc, '3. Website Design, Styling, and Accessibility', 1)
    
    add_heading(doc, 'Accessibility Analysis', 2)
    doc.add_paragraph('The website follows WCAG guidelines for accessibility:')
    doc.add_paragraph('- Semantic Elements: Used <header>, <nav>, <main>, <section>, and <footer> tags.')
    doc.add_paragraph('- Alt Text: All product and logo images include descriptive alternative text for screen readers.')
    doc.add_paragraph('- Color Contrast: High contrast ratios between text and background ensure readability.')
    doc.add_paragraph('- Focus States: Interactive elements have clearly defined hover and focus states.')

    add_heading(doc, 'Responsive Web Design (RWD) Rationale', 2)
    doc.add_paragraph('The website uses CSS Media Queries to adapt the layout between desktop and mobile devices. Flexbox wrap and Grid auto-fit properties enable the product grid to stack vertically on smaller screens, ensuring a seamless experience across all platforms.')
    
    add_heading(doc, 'About Us Page & Contact Form', 2)
    add_image(doc, '09_about_page.png', 'Fig 5: About Us page with semantic elements and CSS layout')
    doc.add_page_break()

    # --- Section 4: Functionality & Storage ---
    add_heading(doc, '4. Website Functionality and Web Storage', 1)
    
    add_heading(doc, 'JavaScript Interaction', 2)
    doc.add_paragraph('JavaScript is utilized for form validation, dynamic alerts, and state management.')
    doc.add_paragraph('- Subscribe Alert: Notifies users upon newsletter signup.')
    doc.add_paragraph('- Add to Cart: Triggers a confirmation alert and updates the cart state.')
    doc.add_paragraph('- Form Submission: Validates user input and provides a "Message Sent" alert.')

    add_heading(doc, 'Web Storage Implementation', 2)
    doc.add_paragraph('Session Storage (sessionStorage): Used to manage the shopping cart. Data persists during the session allowing users to browse without losing items.')
    doc.add_paragraph('Local Storage (localStorage): Used on the About Us page to save contact form metadata, ensuring customer inquiry history is tracked.')

    add_heading(doc, 'Shopping Cart Evidence', 2)
    add_image(doc, '05_cart_modal_items.png', 'Fig 6: Shopping Cart modal displaying items retrieved from sessionStorage')
    
    add_heading(doc, 'Community Corner rendering', 2)
    add_image(doc, '12_community_page.png', 'Fig 7: Community corner showing local testimonials')

    # --- Final Link ---
    doc.add_page_break()
    add_heading(doc, '5. Project Repository Link', 1)
    doc.add_paragraph('The complete project code and documentation is available at the following GitHub repository:')
    p = doc.add_paragraph('https://github.com/PriyanshiYadav007/sophia-touchstone.git')
    p.style.font.color.rgb = None # Set to default or blue if needed

    doc.save(OUTPUT_PATH)
    print(f"Enhanced Final Report compiled at: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
